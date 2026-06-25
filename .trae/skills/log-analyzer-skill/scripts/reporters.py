"""
reporters.py - Multi-format report generators for log-analyzer-skill.
Supports Markdown, JSON, HTML, Word, and PDF report generation.
"""

import json
import html as html_mod
from datetime import datetime
from typing import Dict, Any, Optional

# Optional imports for Word and PDF generation
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False

try:
    from weasyprint import HTML as WP_HTML
    _HAS_WEASYPRINT = True
except ImportError:
    _HAS_WEASYPRINT = False


def generate_markdown(stats: Dict[str, Any], root_cause: Dict[str, Any],
                      file_path: str, duration_sec: float,
                      trend: Optional[Dict[str, Any]] = None,
                      severity: Optional[Dict[str, Any]] = None) -> str:
    """Generate a Markdown analysis report."""
    lines = []
    lines.append("# 📋 日志分析报告")
    lines.append("")
    lines.append(f"- **源文件**: `{file_path}`")
    lines.append(f"- **分析时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"- **耗时**: {duration_sec:.2f}s")
    lines.append("")

    total = stats.get("total_lines", 0)
    errors = stats.get("error_lines", 0)
    rate = stats.get("error_rate", 0)
    tr = stats.get("time_range", {})

    lines.append("## 1. 概览")
    lines.append("")
    lines.append("| 指标 | 值 |")
    lines.append("|------|-----|")
    lines.append(f"| 总行数 | {total:,} |")
    lines.append(f"| 错误行数 | {errors:,} |")
    lines.append(f"| 错误率 | {rate}% |")
    if tr.get("first"):
        lines.append(f"| 时间范围 | {tr.get('first', '')} ~ {tr.get('last', '')} |")

    # Health score
    if severity and severity.get("health_score") is not None:
        hs = severity.get("health_score", 0)
        hs_label = "🟢 良好" if hs >= 80 else "🟡 一般" if hs >= 60 else "🟠 较差" if hs >= 40 else "🔴 严重"
        lines.append(f"| 健康度 | {hs}/100 {hs_label} |")

    lines.append("")

    # Trend analysis
    if trend and trend.get("trend_type") and trend.get("trend_type") != "unknown":
        lines.append("## 1.5 趋势分析")
        lines.append("")
        lines.append(f"- **趋势类型**: {trend.get('trend_type', 'unknown')}")
        lines.append(f"- **描述**: {trend.get('description', '')}")
        lines.append("")

    # Level distribution
    level_icons = {"FATAL": "🔥", "ERROR": "❌", "WARN": "⚠️", "WARNING": "⚠️",
                   "INFO": "ℹ️", "DEBUG": "🔍", "TRACE": "🔬"}
    summary = stats.get("summary", {})
    lines.append("## 2. 日志级别分布")
    lines.append("")
    lines.append("| 级别 | 数量 | 占比 |")
    lines.append("|------|------|------|")
    for level in ["FATAL", "ERROR", "WARN", "WARNING", "INFO", "DEBUG", "TRACE"]:
        count = summary.get(level, 0)
        if count > 0:
            icon = level_icons.get(level, "")
            pct = f"{count / total * 100:.1f}%" if total > 0 else ""
            lines.append(f"| {icon} {level} | {count:,} | {pct} |")
    lines.append("")

    # Error types
    error_types = stats.get("error_types", {})
    if error_types:
        lines.append("## 3. 错误类型分布")
        lines.append("")
        lines.append("| 错误类型 | 次数 |")
        lines.append("|----------|------|")
        for err_type, count in list(error_types.items())[:20]:
            lines.append(f"| {err_type} | {count:,} |")
        lines.append("")

    # Top errors
    top_errors = stats.get("top_errors", [])
    if top_errors:
        n = min(10, len(top_errors))
        lines.append(f"## 4. 高频错误 Top {n}")
        lines.append("")
        lines.append("| # | 错误签名 | 出现次数 |")
        lines.append("|---|---------|---------:|")
        for i, err in enumerate(top_errors[:10], 1):
            msg = err.get("message", "")
            cnt = err.get("count", 0)
            lines.append(f"| {i} | `{_truncate(msg, 80)}` | {cnt} |")
        lines.append("")

    # Root causes with detailed info
    causes = root_cause.get("causes", [])
    if causes:
        lines.append("## 5. 根因分析")
        lines.append("")
        for cause in causes:
            sev = cause.get("severity", "unknown")
            icon = {"critical": "🔥", "high": "⚠️", "medium": "📌", "low": "💡"}.get(sev, "📌")
            lines.append(f"### {icon} {cause.get('name', 'Unknown')}")
            lines.append("")
            lines.append(f"- **类别**: {cause.get('category', '')}")
            lines.append(f"- **严重度**: {sev}")
            lines.append(f"- **出现次数**: {cause.get('count', 0)}")
            if cause.get("cause"):
                lines.append(f"- **原因**: {cause.get('cause', '')}")
            if cause.get("impact"):
                lines.append(f"- **影响**: {cause.get('impact', '')}")
            if cause.get("remediation") and cause["remediation"]:
                lines.append("- **修复措施**:")
                for i, step in enumerate(cause["remediation"], 1):
                    lines.append(f"  {i}. {step}")
            lines.append("")

    # Module distribution
    by_module = stats.get("by_module", {})
    if by_module:
        lines.append("## 6. 模块分布 (错误相关)")
        lines.append("")
        lines.append("| 模块 | 错误数 |")
        lines.append("|------|-------:|")
        for module, count in list(by_module.items())[:20]:
            lines.append(f"| `{module}` | {count} |")
        lines.append("")

    # Hourly trend
    by_hour = stats.get("by_hour", {})
    if by_hour:
        lines.append("## 7. 时间线趋势 (按小时)")
        lines.append("")
        max_count = max(by_hour.values()) if by_hour else 1
        sorted_hours = sorted(by_hour.items())
        for hour, count in sorted_hours[-30:]:
            bar_len = max(1, int(count / max_count * 20))
            bar = "█" * bar_len
            lines.append(f"| {hour} | {bar} {count} |")
        lines.append("")

    # Recommendation
    rec = root_cause.get("recommendation", "")
    if rec:
        lines.append("## 8. 综合建议")
        lines.append("")
        lines.append(f"> {rec}")
        lines.append("")

    return "\n".join(lines)


def generate_json_report(stats: Dict[str, Any], root_cause: Dict[str, Any],
                         file_path: str, duration_sec: float,
                         trend: Optional[Dict[str, Any]] = None,
                         severity: Optional[Dict[str, Any]] = None) -> str:
    """Generate a JSON structured report."""
    report = {
        "report_meta": {
            "tool": "log-analyzer-skill",
            "source_file": file_path,
            "generated_at": datetime.now().isoformat(),
            "duration_sec": round(duration_sec, 2),
        },
        "overview": {
            "total_lines": stats.get("total_lines"),
            "error_lines": stats.get("error_lines"),
            "error_rate": stats.get("error_rate"),
            "time_range": stats.get("time_range", {}),
        },
        "health_score": severity.get("health_score") if severity else None,
        "trend_analysis": trend if trend else None,
        "level_distribution": stats.get("summary", {}),
        "error_types": stats.get("error_types", {}),
        "top_errors": stats.get("top_errors", [])[:20],
        "by_module": stats.get("by_module", {}),
        "by_hour": stats.get("by_hour", {}),
        "root_causes": [
            {
                "id": c.get("id"),
                "name": c.get("name"),
                "severity": c.get("severity"),
                "category": c.get("category"),
                "count": c.get("count"),
                "cause": c.get("cause", ""),
                "impact": c.get("impact", ""),
                "remediation": c.get("remediation", []),
                "suggestion": c.get("suggestion", ""),
            }
            for c in root_cause.get("causes", [])
        ],
        "recommendation": root_cause.get("recommendation", ""),
        "summary": root_cause.get("summary", {}),
    }
    return json.dumps(report, ensure_ascii=False, indent=2)


def generate_html_report(stats: Dict[str, Any], root_cause: Dict[str, Any],
                        file_path: str, duration_sec: float,
                        trend: Optional[Dict[str, Any]] = None,
                        severity: Optional[Dict[str, Any]] = None) -> str:
    """Generate an HTML report with visual formatting."""
    parts = []
    summary = stats.get("summary", {})
    total = stats.get("total_lines", 0)
    errors = stats.get("error_lines", 0)
    rate = stats.get("error_rate", 0)
    tr = stats.get("time_range", {})
    causes = root_cause.get("causes", [])
    top_errors = stats.get("top_errors", [])[:10]
    error_types = stats.get("error_types", {})
    by_module = stats.get("by_module", {})
    by_hour = stats.get("by_hour", {})

    severity_colors = {"critical": "#dc3545", "high": "#fd7e14",
                       "medium": "#ffc107", "low": "#17a2b8"}
    level_icons = {"FATAL": "🔥", "ERROR": "❌", "WARN": "⚠️", "WARNING": "⚠️",
                   "INFO": "ℹ️", "DEBUG": "🔍", "TRACE": "🔬"}

    def esc(s):
        return html_mod.escape(str(s))

    parts.append("""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>日志分析报告</title>
<style>
* { box-sizing: border-box; }
body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 960px; margin: 0 auto; padding: 20px; background: #f8f9fa; color: #333; }
h1 { color: #1a1a2e; border-bottom: 3px solid #0d6efd; padding-bottom: 10px; }
h2 { color: #1a1a2e; margin-top: 30px; }
table { width: 100%; border-collapse: collapse; margin: 10px 0; }
th, td { padding: 8px 12px; text-align: left; border-bottom: 1px solid #dee2e6; }
th { background: #0d6efd; color: white; }
tr:hover { background: #f1f3f5; }
.card { background: white; border-radius: 8px; padding: 20px; margin: 15px 0; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
.stat-row { display: flex; gap: 16px; flex-wrap: wrap; }
.stat-box { flex: 1; min-width: 120px; text-align: center; padding: 16px; background: white; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }
.stat-value { font-size: 28px; font-weight: bold; color: #0d6efd; }
.stat-label { font-size: 13px; color: #6c757d; margin-top: 4px; }
.cause-box { padding: 12px 16px; margin: 10px 0; background: white; border-radius: 4px; box-shadow: 0 1px 3px rgba(0,0,0,0.08); }
.rec-box { background: #e8f4f8; border-radius: 8px; padding: 16px; margin: 10px 0; border-left: 4px solid #17a2b8; }
.footer { text-align: center; color: #6c757d; font-size: 12px; margin-top: 30px; }
</style>
</head>
<body>""")

    parts.append(f"<h1>📋 日志分析报告</h1>")
    parts.append(f"<p><strong>源文件:</strong> <code>{esc(file_path)}</code> &nbsp;|&nbsp;"
                 f" <strong>分析时间:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                 f" &nbsp;|&nbsp; <strong>耗时:</strong> {duration_sec:.2f}s</p>")

    err_color = "#dc3545" if errors > 0 else "#198754"
    rate_color = "#dc3545" if rate > 20 else "#fd7e14" if rate > 5 else "#198754"
    first_ts = tr.get("first", "")
    parts.append(f"""<div class="stat-row">
  <div class="stat-box"><div class="stat-value">{total:,}</div><div class="stat-label">总行数</div></div>
  <div class="stat-box"><div class="stat-value" style="color:{err_color}">{errors:,}</div><div class="stat-label">错误行数</div></div>
  <div class="stat-box"><div class="stat-value" style="color:{rate_color}">{rate}%</div><div class="stat-label">错误率</div></div>""")
    if first_ts:
        parts.append(f"""  <div class="stat-box"><div class="stat-value" style="font-size:14px">{esc(first_ts)}</div><div class="stat-label">起始时间</div></div>""")

    # Health score
    if severity and severity.get("health_score") is not None:
        hs = severity.get("health_score", 0)
        hs_color = "#198754" if hs >= 80 else "#fd7e14" if hs >= 60 else "#dc3545"
        hs_label = "🟢 良好" if hs >= 80 else "🟡 一般" if hs >= 60 else "🟠 较差" if hs >= 40 else "🔴 严重"
        parts.append(f"""  <div class="stat-box"><div class="stat-value" style="color:{hs_color}">{hs}</div><div class="stat-label">健康度 {hs_label}</div></div>""")

    parts.append("</div>")

    # Trend analysis
    if trend and trend.get("trend_type") and trend.get("trend_type") != "unknown":
        trend_icon = {"spike": "📈", "increasing": "📈", "decreasing": "📉", "stable": "➡️"}.get(trend.get("trend_type", ""), "➡️")
        parts.append(f"""<div class="card">
  <h3>{trend_icon} 趋势分析</h3>
  <p>趋势类型: <strong>{trend.get('trend_type', 'unknown')}</strong> — {trend.get('description', '')}</p>
</div>""")

    # Level distribution
    parts.append("""<div class="card"><h2>📊 日志级别分布</h2><table>
<tr><th>级别</th><th>数量</th><th>占比</th></tr>""")
    for level in ["FATAL", "ERROR", "WARN", "WARNING", "INFO", "DEBUG", "TRACE"]:
        cnt = summary.get(level, 0)
        if cnt > 0:
            icon = level_icons.get(level, level)
            pct = f"{cnt / total * 100:.1f}%" if total > 0 else "0%"
            parts.append(f"<tr><td>{icon} {level}</td><td>{cnt:,}</td><td>{pct}</td></tr>")
    parts.append("</table></div>")

    # Error types
    if error_types:
        parts.append("""<div class="card"><h2>🔴 错误类型分布</h2><table>
<tr><th>错误类型</th><th>次数</th></tr>""")
        for err_type, cnt in list(error_types.items())[:15]:
            parts.append(f"<tr><td><code>{esc(err_type)}</code></td><td>{cnt}</td></tr>")
        parts.append("</table></div>")

    # Top errors
    if top_errors:
        parts.append(f"""<div class="card"><h2>🔥 高频错误 Top {min(10, len(top_errors))}</h2><table>
<tr><th>#</th><th>错误签名</th><th>次数</th></tr>""")
        for i, e in enumerate(top_errors[:10], 1):
            msg = esc(e.get("message", ""))[:80]
            cnt = e.get("count", 0)
            parts.append(f"<tr><td>{i}</td><td><code>{msg}</code></td><td>{cnt}</td></tr>")
        parts.append("</table></div>")

    # Root causes
    if causes:
        parts.append("""<div class="card"><h2>🔍 根因分析</h2>""")
        for c in causes[:5]:
            sev = c.get("severity", "medium")
            border_color = severity_colors.get(sev, "#6c757d")
            icon = "🔥" if sev == "critical" else "⚠️" if sev == "high" else "📌"
            name = esc(c.get("name", ""))
            cat = esc(c.get("category", ""))
            cnt = c.get("count", 0)
            sug = esc(c.get("suggestion", ""))
            cause = esc(c.get("cause", ""))
            impact = esc(c.get("impact", ""))
            remediation = c.get("remediation", [])

            parts.append(f"""<div class="cause-box" style="border-left: 4px solid {border_color};">
  <h3 style="margin:0 0 4px">{icon} {name}</h3>
  <p style="margin:4px 0;color:#666;font-size:14px">
    类别: {cat} &nbsp;|&nbsp; 严重度: <strong style="color:{border_color}">{sev}</strong> &nbsp;|&nbsp; 出现 {cnt} 次
  </p>""")

            if cause:
                parts.append(f"""  <p style="margin:4px 0;font-size:14px"><strong>📌 原因:</strong> {cause}</p>""")
            if impact:
                parts.append(f"""  <p style="margin:4px 0;font-size:14px"><strong>⚡ 影响:</strong> {impact}</p>""")
            if remediation and remediation:
                parts.append(f"""  <p style="margin:4px 0;font-size:14px"><strong>🔧 修复措施:</strong></p>
  <ul style="margin:4px 0;padding-left:20px;font-size:14px">""")
                for step in remediation[:5]:
                    parts.append(f"""    <li>{esc(step)}</li>""")
                if len(remediation) > 5:
                    parts.append(f"""    <li><em>... (共 {len(remediation)} 步)</em></li>""")
                parts.append("""  </ul>""")

            parts.append("</div>")
        parts.append("</div>")

    # Hourly trend
    if by_hour:
        parts.append("""<div class="card"><h2>⏱ 时间线趋势</h2><table>
<tr><th>时间</th><th>分布</th></tr>""")
        max_c = max(by_hour.values())
        for hour, cnt in list(sorted(by_hour.items()))[-30:]:
            pct = int(cnt / max_c * 100) if max_c > 0 else 0
            parts.append(f"""<tr><td>{esc(hour)}</td><td><div style="background:#e9ecef;border-radius:4px;overflow:hidden">
<div style="width:{pct}%;background:#0d6efd;color:white;padding:2px 6px;font-size:12px">{cnt}</div></div></td></tr>""")
        parts.append("</table></div>")

    # Module breakdown
    if by_module:
        parts.append("""<div class="card"><h2>📦 模块分布</h2><table>
<tr><th>模块</th><th>错误数</th></tr>""")
        for mod, cnt in list(by_module.items())[:20]:
            parts.append(f"<tr><td><code>{esc(mod)}</code></td><td>{cnt}</td></tr>")
        parts.append("</table></div>")

    # Recommendation
    rec = root_cause.get("recommendation", "")
    if rec:
        parts.append(f"""<div class="rec-box"><h2>💡 综合建议</h2>
<p style="font-size:15px;line-height:1.6">{esc(rec)}</p></div>""")

    parts.append("""<div class="footer">Generated by <strong>log-analyzer-skill</strong> · 数据仅供分析参考</div>
</body></html>""")

    return "\n".join(parts)


def _truncate(text: str, max_len: int) -> str:
    return text[:max_len] + "..." if len(text) > max_len else text


def _severity_color(sev: str) -> str:
    """Return RGB color for severity level."""
    return {
        "critical": (220, 53, 69),
        "high": (253, 126, 20),
        "medium": (255, 193, 7),
        "low": (23, 162, 184),
    }.get(sev, (108, 117, 125))


def generate_word_report(stats: Dict[str, Any], root_cause: Dict[str, Any],
                        file_path: str, duration_sec: float,
                        trend: Optional[Dict[str, Any]] = None,
                        severity: Optional[Dict[str, Any]] = None) -> str:
    """
    Generate a Word (.docx) analysis report.

    Returns empty string if python-docx is not installed.
    """
    if not _HAS_DOCX:
        return ""

    doc = Document()

    # Title
    title = doc.add_heading("日志分析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta info
    meta = doc.add_paragraph()
    meta.add_run("源文件: ").bold = True
    meta.add_run(file_path)
    meta.add_run("\n分析时间: ").bold = True
    meta.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    meta.add_run("\n耗时: ").bold = True
    meta.add_run(f"{duration_sec:.2f}s")

    total = stats.get("total_lines", 0)
    errors = stats.get("error_lines", 0)
    rate = stats.get("error_rate", 0)
    tr = stats.get("time_range", {})
    summary = stats.get("summary", {})
    causes = root_cause.get("causes", [])
    top_errors = stats.get("top_errors", [])[:10]
    error_types = stats.get("error_types", {})
    by_module = stats.get("by_module", {})
    by_hour = stats.get("by_hour", {})

    # Section 1: Overview
    doc.add_heading("1. 概览", 1)
    tbl = doc.add_table(rows=4, cols=2)
    tbl.style = "Table Grid"
    cells = [
        ("总行数", f"{total:,}"),
        ("错误行数", f"{errors:,}"),
        ("错误率", f"{rate}%"),
        ("时间范围", f"{tr.get('first', 'N/A')} ~ {tr.get('last', 'N/A')}"),
    ]
    for i, (k, v) in enumerate(cells):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v

    # Health score
    if severity and severity.get("health_score") is not None:
        doc.add_paragraph()
        hs = severity.get("health_score", 0)
        health_color = "🟢 良好" if hs >= 80 else "🟡 一般" if hs >= 60 else "🟠 较差" if hs >= 40 else "🔴 严重"
        p = doc.add_paragraph()
        p.add_run("健康度评分: ").bold = True
        p.add_run(f"{hs}/100 {health_color}")

    # Trend analysis
    if trend and trend.get("trend_type") != "unknown":
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("趋势分析: ").bold = True
        p.add_run(f"{trend.get('trend_type', 'unknown')} ({trend.get('description', '')})")

    # Section 2: Level distribution
    doc.add_heading("2. 日志级别分布", 1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "级别"
    hdr[1].text = "数量"
    hdr[2].text = "占比"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True

    for level in ["FATAL", "ERROR", "WARN", "WARNING", "INFO", "DEBUG", "TRACE"]:
        cnt = summary.get(level, 0)
        if cnt > 0:
            row = tbl.add_row().cells
            row[0].text = level
            row[1].text = f"{cnt:,}"
            row[2].text = f"{cnt / total * 100:.1f}%" if total > 0 else "0%"

    # Section 3: Error types
    if error_types:
        doc.add_heading("3. 错误类型分布", 1)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "错误类型"
        hdr[1].text = "次数"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
        for err_type, cnt in list(error_types.items())[:20]:
            row = tbl.add_row().cells
            row[0].text = err_type
            row[1].text = f"{cnt:,}"

    # Section 4: Top errors
    if top_errors:
        doc.add_heading(f"4. 高频错误 Top {min(10, len(top_errors))}", 1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "#"
        hdr[1].text = "错误签名"
        hdr[2].text = "次数"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
        for i, err in enumerate(top_errors[:10], 1):
            row = tbl.add_row().cells
            row[0].text = str(i)
            row[1].text = _truncate(err.get("message", ""), 80)
            row[2].text = str(err.get("count", 0))

    # Section 5: Root causes
    if causes:
        doc.add_heading("5. 根因分析", 1)
        for cause in causes[:10]:
            sev = cause.get("severity", "unknown")
            r, g, b = _severity_color(sev)
            color = RGBColor(r, g, b)

            p = doc.add_paragraph()
            run = p.add_run(f"● {cause.get('name', 'Unknown')}")
            run.bold = True
            run.font.color.rgb = color
            run.font.size = Pt(12)

            details = [
                ("类别", cause.get("category", "")),
                ("严重度", sev),
                ("出现次数", str(cause.get("count", 0))),
            ]
            for k, v in details:
                dp = doc.add_paragraph(style="List Bullet")
                dp.add_run(f"{k}: ").bold = True
                dp.add_run(v)

            # Cause, Impact, Remediation
            if cause.get("cause"):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run("原因: ").bold = True
                p.add_run(cause["cause"])

            if cause.get("impact"):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run("影响: ").bold = True
                p.add_run(cause["impact"])

            if cause.get("remediation") and cause["remediation"]:
                p = doc.add_paragraph()
                p.add_run("修复措施:").bold = True
                for step in cause["remediation"]:
                    dp = doc.add_paragraph(style="List Bullet 2")
                    dp.add_run(step)

    # Section 6: Module distribution
    if by_module:
        doc.add_heading("6. 模块分布", 1)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "模块"
        hdr[1].text = "错误数"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
        for mod, cnt in list(by_module.items())[:20]:
            row = tbl.add_row().cells
            row[0].text = mod
            row[1].text = str(cnt)

    # Section 7: Hourly trend
    if by_hour:
        doc.add_heading("7. 时间线趋势", 1)
        sorted_hours = sorted(by_hour.items())
        max_count = max(by_hour.values()) if by_hour else 1
        for hour, count in sorted_hours[-30:]:
            pct = int(count / max_count * 100) if max_count > 0 else 0
            bar = "█" * (pct // 5)
            p = doc.add_paragraph(f"{hour}  {bar} {count}")
            p.style = "List Bullet"

    # Section 8: Recommendation
    rec = root_cause.get("recommendation", "")
    if rec:
        doc.add_heading("8. 综合建议", 1)
        p = doc.add_paragraph(rec)
        p.style = "Quote"

    # Save to file and return path
    output_path = file_path.rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
    if not output_path:
        output_path = "report"
    docx_path = f"/tmp/log_analysis_{output_path}_{int(datetime.now().timestamp())}.docx"
    doc.save(docx_path)
    return docx_path


def generate_pdf_report(stats: Dict[str, Any], root_cause: Dict[str, Any],
                       file_path: str, duration_sec: float,
                       trend: Optional[Dict[str, Any]] = None,
                       severity: Optional[Dict[str, Any]] = None) -> str:
    """
    Generate a PDF analysis report using weasyprint.

    Returns empty string if weasyprint is not installed.
    """
    if not _HAS_WEASYPRINT:
        return ""

    # Generate HTML first, then convert to PDF
    html_content = generate_html_report(stats, root_cause, file_path, duration_sec, trend, severity)

    output_path = file_path.rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
    if not output_path:
        output_path = "report"
    pdf_path = f"/tmp/log_analysis_{output_path}_{int(datetime.now().timestamp())}.pdf"

    wp_html = WP_HTML(string=html_content)
    wp_html.write_pdf(pdf_path)

    return pdf_path


def is_word_available() -> bool:
    """Check if Word report generation is available."""
    return _HAS_DOCX


def is_pdf_available() -> bool:
    """Check if PDF report generation is available."""
    return _HAS_WEASYPRINT

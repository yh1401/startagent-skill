"""
reporters.py - Multi-format report generator for log-analyzer.
Supports: Markdown, JSON, HTML, Word, PDF
"""

import json
import os
from typing import Dict, Any, Optional

try:
    import yaml
    YAML_AVAILABLE = True
except ImportError:
    YAML_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from weasyprint import HTML as WeasyHTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False


def _load_config() -> dict:
    """Load configuration from config.yaml."""
    config_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'config.yaml')
    if YAML_AVAILABLE and os.path.exists(config_path):
        with open(config_path, 'r', encoding='utf-8') as f:
            return yaml.safe_load(f)
    return {}


class ReportGenerator:
    """Multi-format report generator."""

    def __init__(self, output_dir: str = None):
        if output_dir is None:
            config = _load_config()
            output_dir = config.get('report', {}).get('output_dir', './output')
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)

    def generate(self, data: Dict[str, Any], formats: Optional[list] = None) -> Dict[str, str]:
        """
        Generate reports in multiple formats.

        Args:
            data: Analysis data containing stats, root_cause, trends
            formats: List of output formats. Default: ["markdown", "json"]

        Returns:
            Dict mapping format names to file paths
        """
        if formats is None:
            formats = ["markdown", "json"]

        results = {}
        format_handlers = {
            "markdown": self._generate_markdown,
            "json": self._generate_json,
            "html": self._generate_html,
            "word": self._generate_word,
            "pdf": self._generate_pdf,
        }

        for fmt in formats:
            handler = format_handlers.get(fmt)
            if handler:
                try:
                    path = handler(data)
                    results[fmt] = path
                except Exception as e:
                    print(f"Failed to generate {fmt} report: {e}")

        return results

    def _generate_markdown(self, data: Dict[str, Any]) -> str:
        """Generate Markdown report."""
        lines = []
        lines.append("# 日志分析报告")
        lines.append("")
        lines.append(f"**分析时间**: {data.get('analysis_time', '')}")
        lines.append(f"**日志文件**: {data.get('file_path', '')}")
        lines.append("")

        summary = data.get("stats", {}).get("summary", {})
        lines.append("## 1. 概览")
        lines.append("")
        lines.append("| 级别 | 数量 |")
        lines.append("|------|------|")
        for level, count in summary.items():
            lines.append(f"| {level} | {count} |")
        lines.append("")

        error_rate = data.get("stats", {}).get("error_rate", 0)
        total_lines = data.get("stats", {}).get("total_lines", 0)
        lines.append(f"- 总行数: {total_lines}")
        lines.append(f"- 错误率: {error_rate}%")
        lines.append("")

        lines.append("## 2. 严重级别分布")
        lines.append("")
        severity = data.get("severity", {})
        health_score = severity.get("health_score", 0)
        lines.append(f"**健康度评分**: {health_score}/100")
        lines.append("")
        lines.append("| 严重级别 | 数量 | 占比 |")
        lines.append("|----------|------|------|")
        for level in ["critical", "high", "medium", "low"]:
            count = severity.get(level, 0)
            ratio = severity.get("ratios", {}).get(level, 0)
            lines.append(f"| {level} | {count} | {ratio}% |")
        lines.append("")

        lines.append("## 3. 时间趋势分析")
        lines.append("")
        trend = data.get("trend", {})
        lines.append(f"- 趋势类型: {trend.get('trend_type', '')}")
        lines.append(f"- {trend.get('description', '')}")
        lines.append("")

        by_hour = data.get("stats", {}).get("by_hour", {})
        if by_hour:
            lines.append("### 小时分布")
            lines.append("")
            lines.append("| 时间 | 数量 |")
            lines.append("|------|------|")
            for hour, count in by_hour.items():
                lines.append(f"| {hour} | {count} |")
            lines.append("")

        lines.append("## 4. 高频错误")
        lines.append("")
        top_errors = data.get("stats", {}).get("top_errors", [])
        if top_errors:
            lines.append("| 错误消息 | 数量 |")
            lines.append("|----------|------|")
            for err in top_errors[:10]:
                msg = err.get("message", "")[:50] + "..." if len(err.get("message", "")) > 50 else err.get("message", "")
                lines.append(f"| {msg} | {err.get('count', 0)} |")
            lines.append("")
        else:
            lines.append("- 无高频错误")
            lines.append("")

        lines.append("## 5. 按模块统计")
        lines.append("")
        by_module = data.get("stats", {}).get("by_module", {})
        if by_module:
            lines.append("| 模块 | 错误数 |")
            lines.append("|------|--------|")
            for mod, count in list(by_module.items())[:10]:
                lines.append(f"| {mod} | {count} |")
            lines.append("")
        else:
            lines.append("- 无模块统计数据")
            lines.append("")

        lines.append("## 6. 根因分析")
        lines.append("")
        rc = data.get("root_cause", {})
        summary = rc.get("summary", {})
        lines.append(f"**整体评估**: {summary.get('description', '')}")
        lines.append("")

        findings = rc.get("findings", {})
        for severity, issues in findings.items():
            if issues:
                lines.append(f"### {severity.upper()} 级别问题")
                lines.append("")
                for issue in issues:
                    lines.append(f"#### {issue.get('name', '')}")
                    lines.append(f"- **原因**: {issue.get('cause', '')}")
                    lines.append(f"- **影响**: {issue.get('impact', '')}")
                    lines.append("- **建议措施**:")
                    for action in issue.get('remediation', []):
                        lines.append(f"  - {action}")
                    lines.append("")

        lines.append("## 7. 建议")
        lines.append("")
        recommendations = rc.get("recommendations", [])
        if recommendations:
            for rec in recommendations:
                lines.append(f"- **{rec.get('issue', '')}**: {rec.get('cause', '')}")
                for action in rec.get('actions', []):
                    lines.append(f"  - {action}")
                lines.append("")
        else:
            lines.append("- 无特殊建议")

        content = "\n".join(lines)
        filename = f"log_analysis_report.md"
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        return filepath

    def _generate_json(self, data: Dict[str, Any]) -> str:
        """Generate JSON report."""
        filename = f"log_analysis_report.json"
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return filepath

    def _generate_html(self, data: Dict[str, Any]) -> str:
        """Generate HTML report."""
        html = self._build_html(data)
        filename = f"log_analysis_report.html"
        filepath = os.path.join(self.output_dir, filename)
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(html)
        return filepath

    def _build_html(self, data: Dict[str, Any]) -> str:
        """Build HTML content."""
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>日志分析报告</title>
    <style>
        body {{ font-family: 'Microsoft YaHei', sans-serif; margin: 40px; background: #f5f5f5; }}
        .container {{ max-width: 1200px; margin: 0 auto; background: white; padding: 40px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        h1 {{ color: #333; border-bottom: 2px solid #4CAF50; padding-bottom: 10px; }}
        h2 {{ color: #444; margin-top: 30px; }}
        h3 {{ color: #555; }}
        table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
        th, td {{ padding: 12px; text-align: left; border-bottom: 1px solid #ddd; }}
        th {{ background: #f8f9fa; color: #333; }}
        tr:hover {{ background: #f5f5f5; }}
        .badge {{ display: inline-block; padding: 3px 10px; border-radius: 12px; font-size: 12px; color: white; }}
        .badge-critical {{ background: #dc3545; }}
        .badge-high {{ background: #fd7e14; }}
        .badge-medium {{ background: #ffc107; color: #333; }}
        .badge-low {{ background: #17a2b8; }}
        .health-score {{ font-size: 24px; font-weight: bold; color: {'#28a745' if data.get('severity', {}).get('health_score', 0) >= 80 else '#dc3545'}; }}
        .summary-box {{ background: #f8f9fa; padding: 20px; border-radius: 8px; margin: 20px 0; }}
    </style>
</head>
<body>
    <div class="container">
        <h1>日志分析报告</h1>
        <p>分析时间: {data.get('analysis_time', '')}</p>
        <p>日志文件: {data.get('file_path', '')}</p>

        <h2>1. 概览</h2>
        <table>
            <tr><th>级别</th><th>数量</th></tr>
            {''.join(f'<tr><td>{level}</td><td>{count}</td></tr>' for level, count in data.get('stats', {}).get('summary', {}).items())}
        </table>

        <h2>2. 健康度评分</h2>
        <div class="summary-box">
            <span class="health-score">{data.get('severity', {}).get('health_score', 0)}/100</span>
        </div>

        <h2>3. 时间趋势</h2>
        <p>{data.get('trend', {}).get('description', '')}</p>

        <h2>4. 高频错误</h2>
        <table>
            <tr><th>错误消息</th><th>数量</th></tr>
            {''.join(f'<tr><td>{err.get("message", "")[:50]}...</td><td>{err.get("count", 0)}</td></tr>' for err in data.get('stats', {}).get('top_errors', [])[:10])}
        </table>

        <h2>5. 根因分析</h2>
        {self._build_rc_html(data)}
    </div>
</body>
</html>"""

    def _build_rc_html(self, data: Dict[str, Any]) -> str:
        """Build root cause HTML section."""
        rc = data.get("root_cause", {})
        findings = rc.get("findings", {})
        html = ""
        for severity, issues in findings.items():
            if issues:
                html += f"<h3>{severity.upper()} <span class='badge badge-{severity}'>{len(issues)}</span></h3>"
                for issue in issues:
                    html += f"""
                    <h4>{issue.get('name', '')}</h4>
                    <p><strong>原因:</strong> {issue.get('cause', '')}</p>
                    <p><strong>影响:</strong> {issue.get('impact', '')}</p>
                    <p><strong>建议措施:</strong></p>
                    <ul>
                        {''.join(f'<li>{action}</li>' for action in issue.get('remediation', []))}
                    </ul>
                    """
        return html

    def _generate_word(self, data: Dict[str, Any]) -> str:
        """Generate Word report."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install with: pip install python-docx")

        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(10.5)

        heading1 = doc.styles['Heading 1']
        heading1.font.name = 'Microsoft YaHei'
        heading1.font.size = Pt(16)
        heading1.font.bold = True

        heading2 = doc.styles['Heading 2']
        heading2.font.name = 'Microsoft YaHei'
        heading2.font.size = Pt(13)
        heading2.font.bold = True

        doc.add_heading('日志分析报告', level=1)
        doc.add_paragraph(f"分析时间: {data.get('analysis_time', '')}")
        doc.add_paragraph(f"日志文件: {data.get('file_path', '')}")

        summary = data.get("stats", {}).get("summary", {})
        doc.add_heading('1. 概览', level=2)
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '级别'
        hdr_cells[1].text = '数量'
        for level, count in summary.items():
            row_cells = table.add_row().cells
            row_cells[0].text = level
            row_cells[1].text = str(count)

        error_rate = data.get("stats", {}).get("error_rate", 0)
        total_lines = data.get("stats", {}).get("total_lines", 0)
        doc.add_paragraph(f"总行数: {total_lines}")
        doc.add_paragraph(f"错误率: {error_rate}%")

        doc.add_heading('2. 健康度评分', level=2)
        severity = data.get("severity", {})
        doc.add_paragraph(f"健康度评分: {severity.get('health_score', 0)}/100")

        doc.add_heading('3. 时间趋势', level=2)
        trend = data.get("trend", {})
        doc.add_paragraph(f"趋势类型: {trend.get('trend_type', '')}")
        doc.add_paragraph(f"{trend.get('description', '')}")

        doc.add_heading('4. 高频错误', level=2)
        top_errors = data.get("stats", {}).get("top_errors", [])
        if top_errors:
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '错误消息'
            hdr_cells[1].text = '数量'
            for err in top_errors[:10]:
                row_cells = table.add_row().cells
                msg = err.get("message", "")[:50] + "..." if len(err.get("message", "")) > 50 else err.get("message", "")
                row_cells[0].text = msg
                row_cells[1].text = str(err.get("count", 0))

        doc.add_heading('5. 根因分析', level=2)
        rc = data.get("root_cause", {})
        summary = rc.get("summary", {})
        doc.add_paragraph(f"整体评估: {summary.get('description', '')}")

        findings = rc.get("findings", {})
        for severity, issues in findings.items():
            if issues:
                doc.add_heading(f"{severity.upper()} 级别问题", level=3)
                for issue in issues:
                    doc.add_heading(issue.get('name', ''), level=4)
                    doc.add_paragraph(f"原因: {issue.get('cause', '')}")
                    doc.add_paragraph(f"影响: {issue.get('impact', '')}")
                    doc.add_paragraph("建议措施:")
                    for action in issue.get('remediation', []):
                        p = doc.add_paragraph()
                        p.add_run(f"• {action}")

        filename = f"log_analysis_report.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        return filepath

    def _generate_pdf(self, data: Dict[str, Any]) -> str:
        """Generate PDF report."""
        if not WEASYPRINT_AVAILABLE:
            raise ImportError("weasyprint is not installed. Install with: pip install weasyprint")

        html = self._build_html(data)
        filename = f"log_analysis_report.pdf"
        filepath = os.path.join(self.output_dir, filename)
        WeasyHTML(string=html).write_pdf(filepath)
        return filepath


def generate_report(data: Dict[str, Any], formats: Optional[list] = None) -> Dict[str, str]:
    """Convenience function to generate reports."""
    generator = ReportGenerator()
    return generator.generate(data, formats)